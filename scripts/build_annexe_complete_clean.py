from __future__ import annotations

import json
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "livrables_word_mis_en_forme"
EXTERNAL_ANNEXES = Path(r"C:\Users\alban\Documents\À transférer\ECOTEC\Mémoire d'entreprise\Annexes")
CAPTURES_DIR = ROOT / "docs" / "formation_aolink" / "captures"
OUT_WORKSPACE = OUT_DIR / "Annexe_complete_AO_Link.docx"
OUT_EXTERNAL = EXTERNAL_ANNEXES / "Annexe_complete_AO_Link.docx"
MANIFEST = OUT_DIR / "manifest_annexe_complete.json"

LIVRABLES = [
    OUT_DIR / "Livrable_01_Application_web_operationnelle.docx",
    OUT_DIR / "Livrable_02_Documentation_technique.docx",
    OUT_DIR / "Livrable_03_Documentation_de_securite.docx",
    OUT_DIR / "Livrable_04_Scripts_et_fichiers_de_deploiement.docx",
    OUT_DIR / "Livrable_05_Cahier_de_tests_et_resultats.docx",
    OUT_DIR / "Livrable_06_Maintenance_et_exploitation.docx",
    OUT_DIR / "Livrable_07_Annexes_et_documents_de_reference.docx",
]

FORMATION_MD = ROOT / "docs" / "formation_aolink" / "Formation_AO_Link.md"

CAPTURE_ORDER = [
    ("01_connexion_compte_demo.png", "Connexion au compte démo"),
    ("02_liste_projets.png", "Liste des projets"),
    ("03_modal_edition_projet.png", "Édition d'un projet"),
    ("04_modal_partage_projet.png", "Partage d'un projet"),
    ("05_liste_tours.png", "Gestion des tours"),
    ("06_comparaison_tours.png", "Comparaison des tours"),
    ("07_tour_liste_lots.png", "Lots d'un tour"),
    ("08_tour_config_questions.png", "Configuration globale des questions"),
    ("10_lot_donnees_comparatif.png", "Comparatif d'un lot"),
    ("11_lot_donnees_edition.png", "Édition des données"),
    ("12_lot_questions.png", "Questions du lot"),
    ("13_lot_config_questions.png", "Configuration des questions du lot"),
    ("14_parametres.png", "Paramètres et administration"),
]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
    return doc


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_source(doc: Document, source: Path) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(f"Source : {source}")
    run.italic = True
    run.font.size = Pt(8)


def append_table(dst_doc: Document, src_table) -> None:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    is_large = rows > 60 or cols > 10
    if rows > 60 or cols > 10:
        dst_doc.add_paragraph(
            f"Tableau volumineux resume : {rows} lignes x {cols} colonnes. "
            "Apercu limite ci-dessous pour conserver une annexe lisible."
        )
    row_limit = min(rows, 12 if is_large else rows)
    col_limit = min(cols, 6 if is_large else cols)
    if row_limit == 0 or col_limit == 0:
        return
    table = dst_doc.add_table(rows=0, cols=col_limit)
    table.style = "Table Grid"
    for row_index in range(row_limit):
        dst_cells = table.add_row().cells
        for col_index in range(col_limit):
            dst_cells[col_index].text = src_table.cell(row_index, col_index).text
    if is_large:
        dst_doc.add_paragraph("Apercu limite aux 12 premieres lignes et 6 premieres colonnes.")
    dst_doc.add_paragraph("")


def append_paragraph(dst_doc: Document, src_paragraph) -> None:
    text = src_paragraph.text.strip()
    if not text:
        return
    style_name = src_paragraph.style.name if src_paragraph.style else "Normal"
    if style_name not in dst_doc.styles:
        style_name = "Normal"
    p = dst_doc.add_paragraph(style=style_name)
    for run in src_paragraph.runs:
        if not run.text:
            continue
        dst_run = p.add_run(run.text)
        dst_run.bold = run.bold
        dst_run.italic = run.italic
        dst_run.underline = run.underline


def append_docx_content(dst_doc: Document, source: Path) -> None:
    src_doc = Document(source)
    for child in src_doc.element.body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            for paragraph in src_doc.paragraphs:
                if paragraph._element is child:
                    append_paragraph(dst_doc, paragraph)
                    break
        elif tag == "tbl":
            for table in src_doc.tables:
                if table._element is child:
                    append_table(dst_doc, table)
                    break


def add_docx_section(doc: Document, source: Path, index: int, group: str) -> None:
    add_page_break(doc)
    doc.add_heading(f"Annexe {index:02d} - {source.stem.replace('_', ' ')}", level=1)
    p = doc.add_paragraph(f"Groupe : {group}")
    p.runs[0].italic = True
    add_source(doc, source)
    try:
        append_docx_content(doc, source)
    except Exception as exc:
        doc.add_paragraph(f"Impossible d'extraire automatiquement le contenu de ce document : {exc}")


def add_formation_section(doc: Document, index: int) -> None:
    add_page_break(doc)
    doc.add_heading(f"Annexe {index:02d} - Formation utilisateur AO Link", level=1)
    add_source(doc, FORMATION_MD)

    if FORMATION_MD.exists():
        for line in FORMATION_MD.read_text(encoding="utf-8", errors="replace").splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("!["):
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    doc.add_heading("Captures d'ecran de la formation", level=2)
    for filename, caption in CAPTURE_ORDER:
        path = CAPTURES_DIR / filename
        if not path.exists():
            continue
        doc.add_heading(caption, level=3)
        doc.add_picture(str(path), width=Inches(6.8))
        note = doc.add_paragraph(f"Capture : {filename}")
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(8)


def add_xlsx_section(doc: Document, source: Path, index: int) -> None:
    add_page_break(doc)
    doc.add_heading(f"Annexe {index:02d} - {source.stem}", level=1)
    add_source(doc, source)
    doc.add_paragraph(
        "Le planning complet est conserve dans le fichier Excel source. "
        "Pour garder cette annexe lisible, seuls les premiers elements non vides sont repris ci-dessous."
    )
    workbook = load_workbook(source, data_only=True)
    for sheet in workbook.worksheets:
        doc.add_heading(f"Feuille : {sheet.title}", level=2)
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            while values and values[-1] == "":
                values.pop()
            if any(values):
                rows.append(values[:6])
            if len(rows) >= 12:
                break
        if not rows:
            doc.add_paragraph("Feuille vide.")
            continue
        width = max(len(row) for row in rows)
        table = doc.add_table(rows=0, cols=width)
        table.style = "Table Grid"
        for row in rows:
            cells = table.add_row().cells
            for cell_index in range(width):
                cells[cell_index].text = row[cell_index] if cell_index < len(row) else ""
        doc.add_paragraph("Apercu limite aux 12 premieres lignes et 6 premieres colonnes non vides.")


def external_documents() -> tuple[list[Path], list[Path]]:
    docx_files = []
    xlsx_files = []
    if EXTERNAL_ANNEXES.exists():
        for path in sorted(EXTERNAL_ANNEXES.iterdir(), key=lambda p: p.name.lower()):
            if path.name.startswith("~$"):
                continue
            lowered = path.name.lower()
            if lowered.startswith("annexe_complete_ao_link") or lowered.startswith("annexes_ao_link"):
                continue
            if path.suffix.lower() == ".docx":
                docx_files.append(path)
            elif path.suffix.lower() == ".xlsx":
                xlsx_files.append(path)
    return docx_files, xlsx_files


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workspace-only", action="store_true", help="Do not write the copy in the external Annexes folder.")
    parser.add_argument("--external-name", default=OUT_EXTERNAL.name, help="Filename to write in the external Annexes folder.")
    args = parser.parse_args()
    external_output = EXTERNAL_ANNEXES / args.external_name

    doc = setup_document()
    doc.add_paragraph("Annexe complete AO Link", style="Title")
    doc.add_paragraph("Livrables, formation utilisateur et documents annexes du memoire d'entreprise.")
    doc.add_paragraph(f"Document reconstruit le {datetime.now().strftime('%d/%m/%Y a %H:%M')}.")

    external_docx, external_xlsx = external_documents()
    manifest = []
    index = 1

    doc.add_heading("Documents inclus", level=1)
    for path in LIVRABLES:
        if path.exists():
            manifest.append({"index": index, "group": "Livrables Word mis en forme", "path": str(path)})
            doc.add_paragraph(f"{index:02d}. Livrables Word mis en forme - {path.name}")
            index += 1
    if FORMATION_MD.exists():
        manifest.append({"index": index, "group": "Formation utilisateur", "path": str(FORMATION_MD)})
        doc.add_paragraph(f"{index:02d}. Formation utilisateur - {FORMATION_MD.name}")
        index += 1
    for path in external_docx:
        manifest.append({"index": index, "group": "Annexes externes", "path": str(path)})
        doc.add_paragraph(f"{index:02d}. Annexes externes - {path.name}")
        index += 1
    for path in external_xlsx:
        manifest.append({"index": index, "group": "Annexes externes", "path": str(path)})
        doc.add_paragraph(f"{index:02d}. Annexes externes - {path.name}")
        index += 1

    index = 1
    for path in LIVRABLES:
        if path.exists():
            add_docx_section(doc, path, index, "Livrables Word mis en forme")
            index += 1
    if FORMATION_MD.exists():
        add_formation_section(doc, index)
        index += 1
    for path in external_docx:
        add_docx_section(doc, path, index, "Annexes externes")
        index += 1
    for path in external_xlsx:
        add_xlsx_section(doc, path, index)
        index += 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_WORKSPACE)
    if EXTERNAL_ANNEXES.exists() and not args.workspace_only:
        doc.save(external_output)
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Document propre généré : {OUT_WORKSPACE}")
    print(f"Copie memoire : {external_output}")
    print(f"Éléments inclus : {len(manifest)}")


if __name__ == "__main__":
    main()
